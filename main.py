import streamlit as st
import pandas as pd
import numpy as np
import os
import base64
from PIL import Image
import io
import markdown
import zipfile
import textwrap
import json

# Set page config as the FIRST Streamlit command
st.set_page_config(
    page_title="ConvertDOC",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Perform import checks AFTER st.set_page_config()
try:
    import docx  # For python-docx
except ImportError:
    st.error("python-docx module not found. Please install it using: pip install python-docx>=0.8.11")

try:
    import PyPDF2
except ImportError:
    st.error("PyPDF2 module not found. Please install it using: pip install PyPDF2==3.0.1")

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except ImportError:
    st.error("reportlab module not found. Please install it using: pip install reportlab>=4.2.0")

def main():
    # Enhanced CSS for responsiveness, beauty, animations, and vibrant colors
    st.markdown("""
    <style>
    .main {
        background-color: #f5f9ff;
        padding: 20px;
        transition: opacity 0.3s ease-in-out;
    }
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .sidebar .sidebar-content {
        background: linear-gradient(135deg, #2c3e50, #3498db);
        color: white;
        padding: 20px;
        border-radius: 0 10px 10px 0;
        animation: slideInLeft 0.5s ease-out;
    }
    @keyframes slideInLeft {
        from { transform: translateX(-100%); opacity: 0; }
        to { transform: translateX(0); opacity: 1; }
    }
    h1, h2, h3 {
        color: #2c3e50;
        font-family: 'Arial', sans-serif;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    /* Header styling with black background and white text */
    .header {
        background-color: #000000;
        padding: 20px;
        text-align: center;
        border-radius: 5px 5px 0 0;
        animation: fadeIn 0.5s ease-in;
    }
    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }
    .header h1, .header h2 {
        color: white !important;
        margin: 0;
    }
    .stButton>button, .download-button {
        background: linear-gradient(45deg, #2c3e50, #27ae60);
        color: white; /* Ensure button text is white */
        border-radius: 8px;
        padding: 10px 20px;
        border: none;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        font-weight: bold;
        cursor: pointer;
        transition: transform 0.3s ease, opacity 0.3s ease, background 0.3s ease;
        animation: bounceIn 0.5s ease-out;
    }
    @keyframes bounceIn {
        from { transform: scale(0.8); opacity: 0; }
        to { transform: scale(1); opacity: 1; }
    }
    .stButton>button:hover, .download-button:hover {
        background: linear-gradient(45deg, #27ae60, #2c3e50);
        transform: scale(1.05);
        opacity: 0.9;
        box-shadow: 0 6px 8px rgba(0,0,0,0.2);
    }
    .file-converter {
        padding: 20px;
        border-radius: 10px;
        background: linear-gradient(135deg, #ffffff, #e8f5e9);
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        margin-bottom: 20px;
        animation: fadeInUp 0.5s ease-out;
    }
    @keyframes fadeInUp {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    .footer {
        text-align: center;
        margin-top: 50px;
        padding: 20px;
        background: #2c3e50;
        color: white;
        border-radius: 5px;
        box-shadow: 0 -2px 6px rgba(0,0,0,0.1);
        animation: slideInUp 0.5s ease-out;
    }
    @keyframes slideInUp {
        from { transform: translateY(100%); opacity: 0; }
        to { transform: translateY(0); opacity: 1; }
    }
    @media (max-width: 768px) {
        .stApp {
            max-width: 100%;
            padding: 10px;
        }
        .sidebar .sidebar-content {
            border-radius: 0;
        }
        .file-converter {
            padding: 15px;
        }
        .header {
            padding: 10px;
        }
        .stButton>button, .download-button {
            padding: 8px 15px;
            font-size: 14px;
        }
    }
    .developer-section {
        padding: 10px;
        background: #e8f5e9;
        border-radius: 5px;
        margin-top: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        animation: fadeIn 0.5s ease-out;
    }
    .suggestion {
        background-color: #f8f9fa;
        padding: 10px;
        border-radius: 5px;
        margin-top: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        animation: fadeIn 0.5s ease-out;
    }
    </style>
    """, unsafe_allow_html=True)

    # App header with black background, white text, and no logo (centered title/subtitle)
    st.markdown('<div class="header"><h1>ConvertDOC</h1><h2>Developed by Riaz Hussain</h2></div>', unsafe_allow_html=True)

    # Sidebar with navigation and About Developer
    with st.sidebar:
        st.header("Navigation")
        app_mode = st.selectbox("Choose functionality", [
            "File Converter", 
            "Batch Processing", 
            "Text Extraction",
            "About"
        ])

        if st.button("About Developer"):
            with st.expander("About Riaz Hussain", expanded=False):
                st.markdown("""
                <div class="developer-section">
                **Riaz Hussain**  
                I’m a senior student in Quarter 2 at GIAIC (Governor-Sindh Initiative of Artificial Intelligence and Computing). Currently, I’m in my 3rd quarter studying Python and AgentAI while pursuing a Full Stack Developer course.

                **Skills:**
                - HTML & CSS
                - ReactJS
                - Node.js
                - Next.js
                - TailwindCSS
                - TypeScript
                - JavaScript
                - Python

                **Social Media:**
                - [LinkedIn](https://www.linkedin.com/in/riaz-hussain-saifi)
                - [GitHub](https://github.com/Riaz-Hussain-Saifi)
                - [Facebook](https://www.facebook.com/RiazSaifiDeveloper)
                - [WhatsApp](https://wa.me/+923000321640)
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")
        st.write("© 2025 ConvertDOC")
        st.write("All rights reserved")

    # Main content with section-specific details
    if app_mode == "File Converter":
        file_converter()
    elif app_mode == "Batch Processing":
        batch_processing()
    elif app_mode == "Text Extraction":
        text_extraction()
    else:
        about_page()

def file_converter():
    st.header("File Converter")

    st.markdown('<div class="file-converter">', unsafe_allow_html=True)

    st.write("""
    **What You Can Do Here:**
    - **Upload a File:** Select a file (e.g., CSV, Excel, PDF, image, text, JSON) using the uploader above.
    - **Convert Files:** Choose conversion options (e.g., CSV to Excel, PNG to PDF) and click the corresponding button. You’ll see a preview (if applicable) and can download the converted file.
    - **Process Images:** For images, apply filters (Grayscale, Sepia, Invert) or resize them, then convert to PNG, JPG, or PDF.
    - **Data Visualization:** For CSV/Excel, view data as charts (bar, line, scatter) and convert to other formats.
    - **Feedback:** After conversion, a download button will appear for the new file. If there’s an error, you’ll see a message explaining the issue.
    """)

    uploaded_file = st.file_uploader("Upload your file", type=['csv', 'xlsx', 'xls', 'png', 'jpg', 'jpeg', 'pdf', 'txt', 'docx', 'json', 'md'])

    if uploaded_file is not None:
        file_details = {"Filename": uploaded_file.name, "File size": f"{uploaded_file.size/1024:.2f} KB", "File type": uploaded_file.type}
        st.write("### File Details")
        for key, value in file_details.items():
            st.write(f"**{key}:** {value}")

        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        st.write("### Suggestions:")
        st.markdown('<div class="suggestion">', unsafe_allow_html=True)
        if file_extension in ['.csv', '.xlsx', '.xls']:
            st.info("You can convert this data file to CSV, Excel, or visualize it with charts. Try clicking the conversion buttons or chart options below!")
        elif file_extension in ['.png', '.jpg', '.jpeg']:
            st.info("Enhance your image with filters or resizing, then convert to PNG, JPG, or PDF. Click the options below to start!")
        elif file_extension == '.pdf':
            st.info("Extract text from this PDF and download it as a TXT file. Click 'Convert to Text' to proceed!")
        elif file_extension in ['.txt', '.docx', '.md']:
            st.info("Convert this text file to PDF or extract its content. Try the conversion button below!")
        elif file_extension == '.json':
            st.info("Convert this JSON data to CSV for easier analysis. Click 'Convert to CSV' to proceed!")
        st.markdown('</div>', unsafe_allow_html=True)

        if file_extension in ['.csv', '.xlsx', '.xls']:
            convert_data_file(uploaded_file, file_extension)
        elif file_extension in ['.png', '.jpg', 'jpeg']:
            convert_image_file(uploaded_file, file_extension)
        elif file_extension == '.pdf':
            convert_pdf_file(uploaded_file)
        elif file_extension in ['.txt', '.docx', '.md']:
            convert_text_file(uploaded_file, file_extension)
        elif file_extension == '.json':
            convert_json_file(uploaded_file)
        else:
            st.error("Sorry, this file type is not supported for conversion.")

    st.markdown('</div>', unsafe_allow_html=True)

def convert_data_file(uploaded_file, file_extension):
    st.write("### Data File Conversion and Visualization Options")

    try:
        if file_extension == '.csv':
            df = pd.read_csv(uploaded_file)
        else:  # Excel files
            df = pd.read_excel(uploaded_file)

        # Sanitize column names to remove invalid characters (e.g., spaces, colons)
        df.columns = [col.replace(" ", "_").replace(":", "_") for col in df.columns]

        st.dataframe(df.head())

        # Conversion options
        if st.button("Convert to CSV"):
            csv_data = df.to_csv(index=False).encode('utf-8')
            st.success("File converted successfully! You can now download the CSV.")
            download_button(csv_data, "converted_file.csv", "text/csv")

        if st.button("Convert to Excel"):
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df.to_excel(writer, sheet_name='Sheet1', index=False)
            excel_data = output.getvalue()
            st.success("File converted successfully! You can now download the Excel file.")
            download_button(excel_data, "converted_file.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # Chart visualization options
        st.write("### Visualize Data")
        if df.shape[1] >= 2:  # Ensure there are at least 2 columns for charting
            numeric_cols = df.select_dtypes(include=['float64', 'int64']).columns.tolist()
            all_cols = df.columns.tolist()

            if len(numeric_cols) < 1:
                st.warning("No numeric columns available for visualization.")
            else:
                chart_type = st.selectbox("Select chart type", ["Bar Chart", "Line Chart", "Scatter Plot"])
                x_col = st.selectbox("Select X-axis column", all_cols)
                y_col = st.selectbox("Select Y-axis column", numeric_cols)

                if st.button("Generate Chart"):
                    # Ensure data is numeric and properly formatted for Altair
                    if pd.api.types.is_numeric_dtype(df[y_col]):
                        chart_data = df[[x_col, y_col]].dropna()
                        if chart_type == "Bar Chart":
                            st.bar_chart(chart_data.set_index(x_col)[y_col])
                        elif chart_type == "Line Chart":
                            st.line_chart(chart_data.set_index(x_col)[y_col])
                        elif chart_type == "Scatter Plot":
                            st.scatter_chart(chart_data.set_index(x_col)[y_col])
                        st.success("Chart generated successfully!")
                    else:
                        st.error(f"The Y-axis column '{y_col}' must be numeric for charting.")
        else:
            st.warning("Insufficient columns for visualization. Need at least two columns with numeric data.")

    except Exception as e:
        st.error(f"Error processing the file: {e}")

def convert_image_file(uploaded_file, file_extension):
    st.write("### Image File Conversion Options")

    try:
        # Open the image from the uploaded file (BytesIO)
        image_data = uploaded_file.read()
        image = Image.open(io.BytesIO(image_data))
        st.image(image, caption="Uploaded Image", use_container_width=True)

        col1, col2 = st.columns(2)

        with col1:
            st.write("Convert to:")
            to_png = st.button("PNG")
            to_jpg = st.button("JPG")
            to_pdf = st.button("PDF")

        with col2:
            st.write("Image processing:")
            apply_filter = st.selectbox("Apply filter", ["None", "Grayscale", "Sepia", "Invert"])
            resize = st.slider("Resize %", 10, 200, 100)

        # Process image if needed
        if apply_filter != "None" or resize != 100:
            processed_image = process_image(image, apply_filter, resize)
            st.image(processed_image, caption="Processed Image", use_container_width=True)
            image = processed_image

        # Convert image
        if to_png:
            buf = io.BytesIO()
            image.save(buf, format="PNG")
            png_data = buf.getvalue()
            st.success("Image converted to PNG successfully! Download the file below.")
            download_button(png_data, "converted_image.png", "image/png")

        if to_jpg:
            buf = io.BytesIO()
            if image.mode in ("RGBA", "LA"):
                background = Image.new(image.mode[:-1], image.size, (255, 255, 255))
                background.paste(image, image.split()[-1])
                image = background
            image.save(buf, format="JPEG", quality=90)
            jpg_data = buf.getvalue()
            st.success("Image converted to JPG successfully! Download the file below.")
            download_button(jpg_data, "converted_image.jpg", "image/jpeg")

        if to_pdf:
            # Use tempfile for a temporary file path compatible with reportlab
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_file:
                temp_path = temp_file.name
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                image.save(temp_path, format='JPEG', quality=95)

            buf = io.BytesIO()
            pdf = canvas.Canvas(buf, pagesize=letter)
            width, height = letter
            img_width, img_height = image.size

            # Calculate scaling to fit the page, ensuring no distortion
            scale = min(width / img_width, height / img_height) * 0.9
            new_width, new_height = img_width * scale, img_height * scale

            # Use the temporary file path for drawImage
            pdf.drawImage(temp_path, (width - new_width) / 2, (height - new_height) / 2, width=new_width, height=new_height)
            pdf.save()

            pdf_data = buf.getvalue()
            # Clean up temporary file
            os.remove(temp_path)

            st.success("Image converted to PDF successfully! Download the file below.")
            download_button(pdf_data, "converted_image.pdf", "application/pdf")

    except Exception as e:
        st.error(f"Error processing the image: {e}")

def process_image(image, filter_type, resize_percent):
    img = image.copy()

    if resize_percent != 100:
        width, height = img.size
        new_width = int(width * resize_percent / 100)
        new_height = int(height * resize_percent / 100)
        img = img.resize((new_width, new_height), Image.LANCZOS)

    if filter_type == "Grayscale":
        img = img.convert("L")
    elif filter_type == "Sepia":
        img = img.convert("RGB")
        width, height = img.size
        pixels = img.load()
        for py in range(height):
            for px in range(width):
                r, g, b = img.getpixel((px, py))
                tr = int(0.393 * r + 0.769 * g + 0.189 * b)
                tg = int(0.349 * r + 0.686 * g + 0.168 * b)
                tb = int(0.272 * r + 0.534 * g + 0.131 * b)
                pixels[px, py] = (min(tr, 255), min(tg, 255), min(tb, 255))
    elif filter_type == "Invert":
        from PIL import ImageOps
        if img.mode == "RGBA":
            r, g, b, a = img.split()
            rgb_img = Image.merge('RGB', (r, g, b))
            inv_img = ImageOps.invert(rgb_img)
            r2, g2, b2 = inv_img.split()
            img = Image.merge('RGBA', (r2, g2, b2, a))
        else:
            img = ImageOps.invert(img)

    return img

def convert_pdf_file(uploaded_file):
    st.write("### PDF File Conversion Options")

    st.write("""
    **What You Can Do Here:**
    - **Upload a PDF:** Use the uploader to select a PDF file.
    - **Convert to Text:** Click 'Convert to Text' to extract all text from the PDF. You’ll see a preview of the first three pages, and then you can download the full text as a TXT file.
    - **Feedback:** After conversion, a success message will appear, and the download button will be available. If there’s an error, you’ll see a detailed error message.
    """)

    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        num_pages = len(pdf_reader.pages)

        st.write(f"PDF has {num_pages} pages")

        preview_text = ""
        for i in range(min(3, num_pages)):
            page = pdf_reader.pages[i]
            preview_text += page.extract_text() + "\n\n"

        with st.expander("Preview Content (first 3 pages)"):
            st.text(preview_text[:1000] + "..." if len(preview_text) > 1000 else preview_text)

        if st.button("Convert to Text"):
            full_text = ""
            for page in pdf_reader.pages:
                full_text += page.extract_text() + "\n\n"
            st.success("PDF converted to text successfully! Download the TXT file below.")
            download_button(full_text.encode('utf-8'), "converted_file.txt", "text/plain")

    except Exception as e:
        st.error(f"Error processing the PDF: {e}")

def convert_text_file(uploaded_file, file_extension):
    st.write("### Text File Conversion Options")

    st.write("""
    **What You Can Do Here:**
    - **Upload a Text File:** Select a TXT, DOCX, or MD file using the uploader.
    - **Convert to PDF:** Click 'Convert to PDF' to generate a PDF version of the text. You’ll see a preview of the content, and then you can download the PDF.
    - **Feedback:** After conversion, a success message will appear, and the download button will be available. If there’s an error, you’ll see a detailed error message.
    """)

    try:
        if file_extension == '.docx':
            # Use python-docx since docx2txt may not be ideal for Python 3.12
            from docx import Document
            doc = Document(io.BytesIO(uploaded_file.read()))
            text = " ".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == '.md':
            text = uploaded_file.read().decode('utf-8')
            with st.expander("Preview (Rendered Markdown)"):
                st.markdown(text)
        else:  # .txt file
            text = uploaded_file.read().decode('utf-8')

        with st.expander("Preview Content"):
            st.text(text[:1000] + "..." if len(text) > 1000 else text)

        if st.button("Convert to PDF"):
            buf = io.BytesIO()
            pdf = canvas.Canvas(buf, pagesize=letter)
            width, height = letter

            pdf.setFont("Helvetica", 12)
            pdf.drawString(72, height - 72, "Converted Text Document")
            pdf.setFont("Helvetica", 10)

            text_lines = text.split('\n')
            y_position = height - 100
            for line in text_lines:
                if y_position < 72:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 10)
                    y_position = height - 72

                wrapped_lines = textwrap.wrap(line, width=80)
                for wrapped_line in wrapped_lines:
                    pdf.drawString(72, y_position, wrapped_line)
                    y_position -= 15

                y_position -= 5

            pdf.save()
            pdf_data = buf.getvalue()
            st.success("Text converted to PDF successfully! Download the PDF file below.")
            download_button(pdf_data, "converted_file.pdf", "application/pdf")

    except Exception as e:
        st.error(f"Error processing the text file: {e}")

def convert_json_file(uploaded_file):
    st.write("### JSON File Conversion Options")

    st.write("""
    **What You Can Do Here:**
    - **Upload a JSON File:** Use the uploader to select a JSON file.
    - **Convert to CSV:** Click 'Convert to CSV' to transform the JSON data into a CSV format. You’ll see a preview of the JSON structure, and then you can download the CSV.
    - **Feedback:** After conversion, a success message will appear, and the download button will be available. If the JSON structure isn’t suitable (e.g., not a list of objects), you’ll see an error message.
    """)

    try:
        json_data = json.loads(uploaded_file.read().decode('utf-8'))

        with st.expander("Preview JSON"):
            st.json(json_data)

        if st.button("Convert to CSV"):
            if isinstance(json_data, list) and len(json_data) > 0 and isinstance(json_data[0], dict):
                df = pd.json_normalize(json_data)
                csv_data = df.to_csv(index=False).encode('utf-8')
                st.success("JSON converted to CSV successfully! Download the CSV file below.")
                download_button(csv_data, "converted_file.csv", "text/csv")
            else:
                st.error("JSON structure is not suitable for CSV conversion. It should be a list of objects.")

    except Exception as e:
        st.error(f"Error processing the JSON file: {e}")

def batch_processing():
    st.header("Batch Processing")

    st.markdown('<div class="file-converter">', unsafe_allow_html=True)

    st.write("""
    **What You Can Do Here:**
    - **Upload Multiple Files:** Select multiple files of the same type (e.g., CSV, Excel, PDF, images, text) using the uploader.
    - **Process Batch:** Click 'Process Batch' to merge or convert all files into a single ZIP file containing the processed outputs. You’ll see a confirmation of how many files were uploaded, and then you can download the ZIP.
    - **Feedback:** After processing, a success message will appear, and the download button will be available. If there’s an error, you’ll see a detailed error message.
    """)

    uploaded_files = st.file_uploader("Upload multiple files", type=['csv', 'xlsx', 'xls', 'png', 'jpg', 'jpeg', 'pdf', 'txt'], accept_multiple_files=True)

    if uploaded_files:
        st.write(f"Uploaded {len(uploaded_files)} files")

        if st.button("Process Batch"):
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
                for file in uploaded_files:
                    zip_file.writestr(file.name, file.getvalue())
            st.success("Batch processed successfully! Download the ZIP file below.")
            download_button(zip_buffer.getvalue(), "processed_files.zip", "application/zip")

    st.markdown('</div>', unsafe_allow_html=True)

def merge_text_files(files, merge_option, separator):
    merged_text = ""

    for file in files:
        text = file.read().decode('utf-8')
        if merge_option == "Concatenate":
            merged_text += text
        elif merge_option == "Join with newlines":
            merged_text += text + "\n"
        elif merge_option == "Join with separator":
            merged_text += text + separator

    if merge_option == "Join with newlines" and merged_text.endswith("\n"):
        merged_text = merged_text[:-1]
    elif merge_option == "Join with separator" and merged_text.endswith(separator):
        merged_text = merged_text[:-len(separator)]

    return merged_text

def text_extraction():
    st.header("Text Extraction")

    st.markdown('<div class="file-converter">', unsafe_allow_html=True)

    st.write("""
    **What You Can Do Here:**
    - **Upload a Document:** Select a file (PDF, DOCX, TXT, MD, HTML) using the uploader.
    - **Extract Text:** Click 'Download Extracted Text' to extract all text from the document. You’ll see a preview of the extracted text, and then you can download it as a TXT file.
    - **Feedback:** After extraction, a success message will appear, and the download button will be available. If there’s an error, you’ll see a detailed error message.
    """)

    uploaded_file = st.file_uploader("Upload file for text extraction", 
                                   type=['pdf', 'docx', 'txt', 'md', 'html'])

    if uploaded_file:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        try:
            if file_extension == '.pdf':
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

            elif file_extension == '.docx':
                # Use python-docx since docx2txt may not be ideal for Python 3.12
                from docx import Document
                doc = Document(io.BytesIO(uploaded_file.read()))
                text = " ".join([paragraph.text for paragraph in doc.paragraphs])

            elif file_extension in ['.txt', '.md', '.html']:
                text = uploaded_file.read().decode('utf-8')

            st.text_area("Extracted Text", text, height=400)

            if st.button("Download Extracted Text"):
                st.success("Text extracted successfully! Download the TXT file below.")
                download_button(text.encode('utf-8'), "extracted_text.txt", "text/plain")

        except Exception as e:
            st.error(f"Error extracting text: {e}")

    st.markdown('</div>', unsafe_allow_html=True)

def about_page():
    st.header("About ConvertDOC")

    st.markdown('<div class="file-converter">', unsafe_allow_html=True)

    st.write("""
    **ConvertDOC** is a powerful, user-friendly application designed to simplify your file management and conversion tasks. It offers the following features:

    - **File Converter:** Convert between various file formats like CSV, Excel, PDF, images, text files, and JSON with ease. Includes data visualization and image processing options.
    - **Batch Processing:** Handle multiple files simultaneously, merging or converting them into a single output for efficiency.
    - **Text Extraction:** Extract text from documents for easy access or further use, with the option to download as TXT files.

    This tool is ideal for developers, data analysts, and anyone needing efficient file conversion and processing solutions.
    """)

    st.write("© 2025 ConvertDOC. All rights reserved.")

    st.markdown('</div>', unsafe_allow_html=True)

def download_button(data, file_name, mime_type):
    b64 = base64.b64encode(data).decode()
    href = f'<a class="download-button" href="data:{mime_type};base64,{b64}" download="{file_name}">{file_name}</a>'
    st.markdown(href, unsafe_allow_html=True)

if __name__ == "__main__":
    main()